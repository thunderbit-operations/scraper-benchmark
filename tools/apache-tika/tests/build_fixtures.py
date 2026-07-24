#!/usr/bin/env python3
"""build_fixtures.py — SINGLE SOURCE OF TRUTH for the Apache Tika evaluation fixtures.

Three fixture families, all emitted here so rendered bytes and intended labels can never
drift:

  A. CONTENT + METADATA (multi-carrier).  One logical document rendered across every
     Tika-native, dependency-free carrier: html / md / txt / docx / pdf / rtf / odt (and
     a csv/txt pair for the table doc).  Every logical block carries a UNIQUE sentinel
     token embedded in REAL prose; known author/title/created metadata is embedded where
     the carrier supports it.  Downstream matching is exact sentinel substring membership
     (never fuzzy).  This drives H2 (content-extraction recall) + H3 (metadata behavior).

  B. ADVERSARIAL MIME (extension x filename matrix).  Real files of known true type, each
     also copied under (i) a LYING extension (a different type's extension) and (ii) NO
     extension.  run_tika.py detects each with the filename present AND via stdin (no
     filename).  Ground truth = the true media type.  Drives H1 (magic vs extension).

  C. ADVERSARIAL ROBUSTNESS.  empty file, truncated binaries (first N bytes of pdf/docx),
     and a charset case (UTF-8 non-ASCII body) — for graceful-degradation + charset
     detection.  Drives H4.

Determinism: pure string assembly + fixed content for the text carriers.  The only
non-deterministic bytes live inside the .docx / .odt zips and the .pdf (internal
timestamps); those are gitignored and regenerated, never asserted byte-wise — only the
EXTRACTED text / metadata / detected type are asserted.

ground_truth.json records, per logical doc: the ordered blocks {id,type,sentinel,text}
and the embedded metadata {author,title,created} with per-field applicability per format.
mime_truth.json records, per adversarial artifact: true_media_type + the extension/filename
condition.  robustness_truth.json records the C-family expectations.
"""
from __future__ import annotations

import html as htmllib
import json
import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from odf.opendocument import OpenDocumentText
from odf import text as odftext, meta as odfmeta, dc as odfdc
from odf.namespaces import METANS

HERE = Path(__file__).resolve().parent
FIX = HERE / "fixtures"
FIX.mkdir(parents=True, exist_ok=True)

# ---- known embedded metadata (ground truth for H3) --------------------------------------
META_AUTHOR = "Ada Lovelace zzmauth"
META_TITLE = "Analytical Engine Report zzmtitle"
META_CREATED = datetime(2021, 3, 15, 9, 30, 0)
META_CREATED_ISO = "2021-03-15T09:30:00Z"


# =========================================================================================
# A. Content + metadata: logical documents (sentinel-annotated real prose)
# =========================================================================================
class Block:
    def __init__(self, id, intent, sentinel, text, kind, level=1, extra=None):
        self.id, self.intent, self.sentinel = id, intent, sentinel
        self.text, self.kind, self.level, self.extra = text, kind, level, extra


def gt_of(b: Block) -> dict:
    return {"id": b.id, "type": b.intent, "sentinel": b.sentinel, "text": b.text}


def doc_canonical() -> list[Block]:
    # Real English prose (real verbs), one unique sentinel per block. Mirrors the sibling
    # unstructured/docling fixture shape so the content-recall numbers are same-testbed
    # comparable — but here we score CONTENT SURVIVAL, not element-type classification
    # (Tika emits flat text + metadata, it does not label Title/NarrativeText/ListItem).
    return [
        Block("c-t1", "Title", "zztitle1", "Quarterly Operations Review zztitle1", "title", 1),
        Block("c-n1", "NarrativeText", "zznarr1",
              "Our team zznarr1 evaluated fourteen extraction tools during the pilot and recorded every result.", "narr"),
        Block("c-n2", "NarrativeText", "zznarr2",
              "The engineers zznarr2 compared latency and memory across each candidate before writing the recommendation.", "narr"),
        Block("c-t2", "Title", "zztitle2", "Key Findings zztitle2", "title", 2),
        Block("c-b1", "ListItem", "zzitem1", "First zzitem1 the parser handled malformed markup without crashing", "bullet"),
        Block("c-b2", "ListItem", "zzitem2", "Second zzitem2 it preserved the original list ordering throughout", "bullet"),
        Block("c-b3", "ListItem", "zzitem3", "Third zzitem3 it exported clean readable plain text reliably", "bullet"),
        Block("c-t3", "Title", "zztitle3", "Setup Steps zztitle3", "title", 2),
        Block("c-o1", "ListItem", "zznum1", "Install zznum1 the project dependencies from the lockfile", "number"),
        Block("c-o2", "ListItem", "zznum2", "Run zznum2 the extractor against the sample corpus", "number"),
        Block("c-n3", "NarrativeText", "zznarr3",
              "In summary zznarr3 the tool performed consistently across every evaluated document format.", "narr"),
    ]


def doc_table() -> list[Block]:
    return [
        Block("t-t1", "Title", "zztbltitle", "Benchmark Results zztbltitle", "title", 1),
        Block("t-n1", "NarrativeText", "zztblnarr",
              "The table zztblnarr below summarizes the measured throughput for every tool we tested.", "narr"),
        Block("t-tab", "Table", "zztblcell", "table with sentinel cells", "table",
              extra=[["Tool", "Throughput"], ["zztblcell_alpha", "120"], ["zztblcell_beta", "95"]]),
    ]


# ---- renderers: same block list -> many carriers ----------------------------------------
def render_html(blocks, with_meta=True) -> str:
    head = ["<!DOCTYPE html><html lang=\"en\"><head><meta charset=\"utf-8\">"]
    if with_meta:
        head.append(f"<title>{htmllib.escape(META_TITLE)}</title>")
        head.append(f"<meta name=\"author\" content=\"{htmllib.escape(META_AUTHOR)}\">")
    else:
        head.append("<title>fixture</title>")
    head.append("</head><body>")
    parts = head
    i = 0
    while i < len(blocks):
        b = blocks[i]
        if b.kind == "title":
            parts.append(f"<h{b.level}>{htmllib.escape(b.text)}</h{b.level}>"); i += 1
        elif b.kind in ("narr", "raw"):
            parts.append(f"<p>{htmllib.escape(b.text)}</p>"); i += 1
        elif b.kind in ("bullet", "number"):
            tag = "ul" if b.kind == "bullet" else "ol"
            parts.append(f"<{tag}>")
            while i < len(blocks) and blocks[i].kind == b.kind:
                parts.append(f"<li>{htmllib.escape(blocks[i].text)}</li>"); i += 1
            parts.append(f"</{tag}>")
        elif b.kind == "table":
            parts.append("<table>")
            for r, row in enumerate(b.extra):
                cell = "th" if r == 0 else "td"
                parts.append("<tr>" + "".join(f"<{cell}>{htmllib.escape(c)}</{cell}>" for c in row) + "</tr>")
            parts.append("</table>"); i += 1
        else:
            i += 1
    parts.append("</body></html>")
    return "\n".join(parts) + "\n"


def render_md(blocks) -> str:
    lines = []
    i = 0
    while i < len(blocks):
        b = blocks[i]
        if b.kind == "title":
            lines += [("#" * b.level) + " " + b.text, ""]; i += 1
        elif b.kind in ("narr", "raw"):
            lines += [b.text, ""]; i += 1
        elif b.kind in ("bullet", "number"):
            if lines and lines[-1] != "":
                lines.append("")
            n = 1
            while i < len(blocks) and blocks[i].kind == b.kind:
                lines.append(("- " if b.kind == "bullet" else f"{n}. ") + blocks[i].text); n += 1; i += 1
            lines.append("")
        elif b.kind == "table":
            if lines and lines[-1] != "":
                lines.append("")
            rows = b.extra
            lines.append("| " + " | ".join(rows[0]) + " |")
            lines.append("|" + "|".join([" --- "] * len(rows[0])) + "|")
            for row in rows[1:]:
                lines.append("| " + " | ".join(row) + " |")
            lines.append(""); i += 1
        else:
            i += 1
    return "\n".join(lines) + "\n"


def render_txt(blocks) -> str:
    lines = []
    for b in blocks:
        if b.kind in ("title", "narr", "raw"):
            lines += [b.text, ""]
        elif b.kind == "bullet":
            lines += ["- " + b.text, ""]
        elif b.kind == "number":
            lines += ["1. " + b.text, ""]
        elif b.kind == "table":
            for row in b.extra:
                lines.append("  ".join(row))
            lines.append("")
    return "\n".join(lines) + "\n"


def render_csv(blocks) -> str:
    # only the table doc's grid, as real CSV (regular column count -> exercises Tika's
    # statistical CSV detector). Non-table blocks are dropped: a CSV carries only the grid.
    for b in blocks:
        if b.kind == "table":
            return "\n".join(",".join(row) for row in b.extra) + "\n"
    return ""


def render_xml(blocks) -> str:
    parts = ["<?xml version=\"1.0\" encoding=\"UTF-8\"?>", "<document>"]
    for b in blocks:
        if b.kind == "table":
            parts.append("<table>")
            for row in b.extra:
                parts.append("<row>" + "".join(f"<cell>{htmllib.escape(c)}</cell>" for c in row) + "</row>")
            parts.append("</table>")
        else:
            parts.append(f"<block type=\"{b.intent}\">{htmllib.escape(b.text)}</block>")
    parts.append("</document>")
    return "\n".join(parts) + "\n"


def _rtf_escape(s: str) -> str:
    return s.replace("\\", "\\\\").replace("{", "\\{").replace("}", "\\}")


def render_rtf(blocks) -> str:
    # minimal but valid RTF: {\rtf1 ... \par separated paragraphs}. Bullets get a "- "
    # prefix (RTF list structure is optional; we test text survival, not list semantics).
    parts = ["{\\rtf1\\ansi\\deff0"]
    for b in blocks:
        if b.kind == "table":
            for row in b.extra:
                parts.append(_rtf_escape("  ".join(row)) + "\\par")
        elif b.kind == "bullet":
            parts.append(_rtf_escape("- " + b.text) + "\\par")
        elif b.kind == "number":
            parts.append(_rtf_escape("1. " + b.text) + "\\par")
        else:
            parts.append(_rtf_escape(b.text) + "\\par")
    parts.append("}")
    return "\n".join(parts) + "\n"


def render_docx(blocks, path: Path, with_meta=True) -> None:
    d = Document()
    if with_meta:
        d.core_properties.author = META_AUTHOR
        d.core_properties.title = META_TITLE
        d.core_properties.created = META_CREATED
    for b in blocks:
        if b.kind == "title":
            d.add_heading(b.text, level=b.level)
        elif b.kind in ("narr", "raw"):
            d.add_paragraph(b.text)
        elif b.kind == "bullet":
            d.add_paragraph(b.text, style="List Bullet")
        elif b.kind == "number":
            d.add_paragraph(b.text, style="List Number")
        elif b.kind == "table":
            rows = b.extra
            t = d.add_table(rows=len(rows), cols=len(rows[0]))
            for r, row in enumerate(rows):
                for c, val in enumerate(row):
                    t.rows[r].cells[c].text = val
    d.save(str(path))


def render_pdf(blocks, path: Path, with_meta=True) -> None:
    c = canvas.Canvas(str(path), pagesize=letter)
    if with_meta:
        c.setAuthor(META_AUTHOR)
        c.setTitle(META_TITLE)
    y = 740
    for b in blocks:
        if b.kind == "table":
            for row in b.extra:
                c.drawString(72, y, "  ".join(row)); y -= 20
        elif b.kind == "bullet":
            c.drawString(72, y, "- " + b.text); y -= 20
        elif b.kind == "number":
            c.drawString(72, y, "1. " + b.text); y -= 20
        else:
            c.drawString(72, y, b.text); y -= 20
        if y < 72:
            c.showPage(); y = 740
    c.save()


def render_odt(blocks, path: Path, with_meta=True) -> None:
    doc = OpenDocumentText()
    if with_meta:
        doc.meta.addElement(odfdc.Creator(text=META_AUTHOR))
        doc.meta.addElement(odfdc.Title(text=META_TITLE))
        cd = odfmeta.CreationDate(text=META_CREATED.strftime("%Y-%m-%dT%H:%M:%S"))
        doc.meta.addElement(cd)
    for b in blocks:
        if b.kind == "table":
            for row in b.extra:
                doc.text.addElement(odftext.P(text="  ".join(row)))
        elif b.kind == "bullet":
            doc.text.addElement(odftext.P(text="- " + b.text))
        elif b.kind == "number":
            doc.text.addElement(odftext.P(text="1. " + b.text))
        else:
            doc.text.addElement(odftext.P(text=b.text))
    doc.save(str(path))


# ---- metadata applicability per format (ground truth for H3) ----------------------------
# Which of {author,title,created} we actually EMBED in each carrier (so a miss is a real
# miss, not us asking for a field we never wrote). txt/md/csv/xml carry NO document metadata.
META_APPLICABILITY = {
    "html": {"author": True, "title": True, "created": False},
    "docx": {"author": True, "title": True, "created": True},
    "pdf":  {"author": True, "title": True, "created": False},   # reportlab sets created=now, not our known value -> presence-only, not scored exact
    "odt":  {"author": True, "title": True, "created": True},
    "rtf":  {"author": False, "title": False, "created": False},
    "txt":  {"author": False, "title": False, "created": False},
    "md":   {"author": False, "title": False, "created": False},
    "csv":  {"author": False, "title": False, "created": False},
    "xml":  {"author": False, "title": False, "created": False},
}


def emit_content_fixtures() -> dict:
    docs = {}

    canon = doc_canonical()
    canon_formats = ["html", "md", "txt", "docx", "pdf", "rtf", "odt", "xml"]
    (FIX / "canonical.html").write_text(render_html(canon), encoding="utf-8")
    (FIX / "canonical.md").write_text(render_md(canon), encoding="utf-8")
    (FIX / "canonical.txt").write_text(render_txt(canon), encoding="utf-8")
    (FIX / "canonical.xml").write_text(render_xml(canon), encoding="utf-8")
    (FIX / "canonical.rtf").write_text(render_rtf(canon), encoding="utf-8")
    render_docx(canon, FIX / "canonical.docx")
    render_pdf(canon, FIX / "canonical.pdf")
    render_odt(canon, FIX / "canonical.odt")
    canon_ids = [b.id for b in canon]
    docs["canonical"] = {
        "name": "canonical", "formats": canon_formats,
        "blocks": [gt_of(b) for b in canon],
        # every canonical carrier renders every block (a full document in each format)
        "blocks_by_format": {f: canon_ids for f in canon_formats},
        "metadata": {"author": META_AUTHOR, "title": META_TITLE, "created": META_CREATED_ISO},
        "metadata_applicability": {f: META_APPLICABILITY[f] for f in canon_formats},
    }

    tbl = doc_table()
    tbl_formats = ["html", "md", "txt", "docx", "csv", "xml"]
    (FIX / "table.html").write_text(render_html(tbl, with_meta=False), encoding="utf-8")
    (FIX / "table.md").write_text(render_md(tbl), encoding="utf-8")
    (FIX / "table.txt").write_text(render_txt(tbl), encoding="utf-8")
    (FIX / "table.csv").write_text(render_csv(tbl), encoding="utf-8")
    (FIX / "table.xml").write_text(render_xml(tbl), encoding="utf-8")
    render_docx(tbl, FIX / "table.docx", with_meta=False)
    tbl_all = [b.id for b in tbl]
    tbl_tab_only = [b.id for b in tbl if b.kind == "table"]
    docs["table"] = {
        "name": "table", "formats": tbl_formats,
        "blocks": [gt_of(b) for b in tbl],
        # a CSV carries ONLY the grid — title/narrative blocks are not part of a CSV
        # rendering, so recall for csv is scored against the table block alone (measuring
        # Tika's extraction, not the fixture's coverage). All other carriers render all 3.
        "blocks_by_format": {f: (tbl_tab_only if f == "csv" else tbl_all) for f in tbl_formats},
        "metadata": None,
        "metadata_applicability": {f: {"author": False, "title": False, "created": False} for f in tbl_formats},
    }
    return docs


# =========================================================================================
# B. Adversarial MIME matrix: true type x {correct ext, lying ext, no ext}
# =========================================================================================
# Each entry: a source fixture file (a real, well-formed file of TRUE_TYPE) + the true media
# type Tika should ideally report. run_tika.py detects each variant with the filename (glob)
# AND from a filename-less stdin (content/magic only).
MIME_SOURCES = [
    ("canonical.pdf",  "application/pdf",  "pdf",  True),   # %PDF magic
    ("canonical.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", "docx", True),  # zip+[Content_Types]
    ("canonical.rtf",  "application/rtf",  "rtf",  True),   # {\rtf magic
    ("canonical.html", "text/html",        "html", True),   # <!DOCTYPE/<html sniff
    ("canonical.xml",  "application/xml",   "xml",  True),   # <?xml root
    ("canonical.md",   "text/markdown",    "md",   False),  # NO magic -> extension-only
    ("canonical.txt",  "text/plain",       "txt",  False),  # NO magic
    ("table.csv",      "text/csv",         "csv",  False),  # statistical detector (borderline)
]
# lying-extension target: map each true type to a DIFFERENT type's extension to create a lie.
LYING_EXT = {"pdf": "txt", "docx": "jpg", "rtf": "html", "html": "csv",
             "xml": "txt", "md": "pdf", "txt": "pdf", "csv": "txt"}


def emit_mime_fixtures() -> list[dict]:
    mdir = FIX / "mime"
    mdir.mkdir(exist_ok=True)
    entries = []
    for src, true_type, tag, has_magic in MIME_SOURCES:
        srcpath = FIX / src
        base = f"m_{tag}"
        correct = mdir / f"{base}.{src.split('.')[-1]}"
        lying = mdir / f"{base}_lying.{LYING_EXT[tag]}"
        noext = mdir / f"{base}_noext"
        shutil.copyfile(srcpath, correct)
        shutil.copyfile(srcpath, lying)
        shutil.copyfile(srcpath, noext)
        entries.append({
            "tag": tag, "true_media_type": true_type, "has_magic": has_magic,
            "variants": {
                "correct_ext": correct.name,
                "lying_ext": lying.name,
                "no_ext": noext.name,
            },
            "lying_ext_is": LYING_EXT[tag],
        })
    return entries


# =========================================================================================
# C. Adversarial robustness: empty, truncated, non-ASCII charset
# =========================================================================================
def emit_robustness_fixtures() -> list[dict]:
    rdir = FIX / "robust"
    rdir.mkdir(exist_ok=True)
    entries = []

    (rdir / "empty.txt").write_bytes(b"")
    entries.append({"name": "empty.txt", "case": "empty",
                    "expect_no_crash": True, "true_media_type": None})

    # truncated PDF: first 512 bytes of the real pdf (valid %PDF header, broken body)
    pdf_bytes = (FIX / "canonical.pdf").read_bytes()
    (rdir / "truncated.pdf").write_bytes(pdf_bytes[:512])
    entries.append({"name": "truncated.pdf", "case": "truncated_pdf",
                    "expect_no_crash": True, "true_media_type": "application/pdf"})

    # truncated DOCX: first 1024 bytes of the zip (PK header, broken central directory)
    docx_bytes = (FIX / "canonical.docx").read_bytes()
    (rdir / "truncated.docx").write_bytes(docx_bytes[:1024])
    entries.append({"name": "truncated.docx", "case": "truncated_docx",
                    "expect_no_crash": True, "true_media_type_family": "zip/ooxml"})

    # non-ASCII UTF-8 body (no BOM, no charset declaration) -> exercises encoding detector.
    # Known sentinel + multibyte chars whose correct decoding proves UTF-8 was chosen.
    utf8_body = "Charset zzutf8 probe: café — naïve — Ω ≈ π — 日本語テスト.\n"
    (rdir / "utf8_nobom.txt").write_bytes(utf8_body.encode("utf-8"))
    entries.append({"name": "utf8_nobom.txt", "case": "utf8_no_bom",
                    "expect_no_crash": True, "true_media_type": "text/plain",
                    "sentinel": "zzutf8", "must_contain": "日本語テスト",
                    "true_charset_family": "UTF-8"})

    return entries


def main() -> int:
    content = emit_content_fixtures()
    (FIX / "ground_truth.json").write_text(
        json.dumps(content, indent=2, ensure_ascii=False) + "\n", encoding="utf-8")

    mime = emit_mime_fixtures()
    (FIX / "mime_truth.json").write_text(
        json.dumps({"entries": mime,
                    "note": "true_media_type is the ideal report; magic-less text formats "
                            "(md/txt/csv) legitimately fall back to text/plain when the "
                            "extension is absent or lies — that is the measured behavior, "
                            "not a bug."}, indent=2, ensure_ascii=False) + "\n", encoding="utf-8")

    robust = emit_robustness_fixtures()
    (FIX / "robustness_truth.json").write_text(
        json.dumps({"entries": robust}, indent=2, ensure_ascii=False) + "\n", encoding="utf-8")

    n_docs = len(content)
    n_blocks = sum(len(d["blocks"]) for d in content.values())
    n_carriers = sum(len(d["formats"]) for d in content.values())
    print(f"content: {n_docs} logical docs, {n_blocks} labeled blocks, {n_carriers} carrier renderings")
    print(f"mime:    {len(mime)} true types x 3 ext-conditions x 2 filename-conditions")
    print(f"robust:  {len(robust)} adversarial-degradation cases")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
