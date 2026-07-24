#!/usr/bin/env python3
"""build_fixtures.py — SINGLE SOURCE OF TRUTH for the cross-format fixtures.

For each LOGICAL document it emits up to four carrier renderings under tests/fixtures/:
  <doc>.html  <doc>.md  <doc>.txt  <doc>.docx      (docx via python-docx)
and one shared tests/fixtures/ground_truth.json, so the rendered bytes and the intended
element-type labels can NEVER drift.

Each logical block carries a UNIQUE sentinel token (e.g. "zztitle1") embedded inside
REAL prose (real English with real verbs, so the heuristic classifier is exercised
fairly — a narrative paragraph must actually contain a verb to be judged NarrativeText).
Downstream matching is exact sentinel substring membership, never fuzzy.

Ground truth per block: { id, type, sentinel, text, note? }.
  type ∈ {Title, NarrativeText, ListItem, Table}   (the author's INTENT)

Determinism: pure string assembly, fixed content, no randomness. The only non-deterministic
bytes are inside the .docx zip (internal timestamps); extraction is unaffected and the
.docx is gitignored + regenerated, never asserted byte-wise.
"""
from __future__ import annotations

import html as htmllib
import json
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

HERE = Path(__file__).resolve().parent
FIX = HERE / "fixtures"
FIX.mkdir(parents=True, exist_ok=True)


# ---------------------------------------------------------------------------
# A block is the atomic labeled unit. `text` is the natural-language content
# (already contains the sentinel token). `kind` drives how each carrier renders it.
#   kind: "title" (heading), "narr" (body paragraph), "bullet", "number",
#         "table" (rows: list[list[str]] in .extra), "raw" (verbatim body line)
# `level` = heading level for titles (1/2). `intent` = the ground-truth type.
# ---------------------------------------------------------------------------
class Block:
    def __init__(self, id, intent, sentinel, text, kind, level=1, extra=None, note=None):
        self.id = id
        self.intent = intent
        self.sentinel = sentinel
        self.text = text
        self.kind = kind
        self.level = level
        self.extra = extra
        self.note = note


def gt_of(block: Block) -> dict:
    d = {"id": block.id, "type": block.intent, "sentinel": block.sentinel, "text": block.text}
    if block.note:
        d["note"] = block.note
    return d


# ---------------------------------------------------------------------------
# Renderers: same block list -> four carriers. Grouping consecutive bullets /
# numbers / table rows into a single list/table construct is handled here.
# ---------------------------------------------------------------------------
def render_html(blocks: list[Block]) -> str:
    parts = ["<!DOCTYPE html><html lang=\"en\"><head><meta charset=\"utf-8\">",
             "<title>fixture</title></head><body>"]
    i = 0
    while i < len(blocks):
        b = blocks[i]
        if b.kind == "title":
            parts.append(f"<h{b.level}>{htmllib.escape(b.text)}</h{b.level}>")
            i += 1
        elif b.kind == "narr":
            parts.append(f"<p>{htmllib.escape(b.text)}</p>")
            i += 1
        elif b.kind == "raw":
            parts.append(f"<p>{htmllib.escape(b.text)}</p>")
            i += 1
        elif b.kind in ("bullet", "number"):
            tag = "ul" if b.kind == "bullet" else "ol"
            parts.append(f"<{tag}>")
            while i < len(blocks) and blocks[i].kind == b.kind:
                parts.append(f"<li>{htmllib.escape(blocks[i].text)}</li>")
                i += 1
            parts.append(f"</{tag}>")
        elif b.kind == "table":
            parts.append("<table>")
            for r, row in enumerate(b.extra):
                cell = "th" if r == 0 else "td"
                parts.append("<tr>" + "".join(f"<{cell}>{htmllib.escape(c)}</{cell}>" for c in row) + "</tr>")
            parts.append("</table>")
            i += 1
        else:
            i += 1
    parts.append("</body></html>")
    return "\n".join(parts) + "\n"


def render_md(blocks: list[Block], blank_before_list: bool = True) -> str:
    lines = []
    i = 0
    while i < len(blocks):
        b = blocks[i]
        if b.kind == "title":
            lines.append(("#" * b.level) + " " + b.text)
            lines.append("")
            i += 1
        elif b.kind in ("narr", "raw"):
            lines.append(b.text)
            lines.append("")
            i += 1
        elif b.kind in ("bullet", "number"):
            if blank_before_list and lines and lines[-1] != "":
                lines.append("")
            n = 1
            while i < len(blocks) and blocks[i].kind == b.kind:
                prefix = "- " if b.kind == "bullet" else f"{n}. "
                lines.append(prefix + blocks[i].text)
                n += 1
                i += 1
            lines.append("")
        elif b.kind == "table":
            if lines and lines[-1] != "":
                lines.append("")
            rows = b.extra
            lines.append("| " + " | ".join(rows[0]) + " |")
            lines.append("|" + "|".join([" --- "] * len(rows[0])) + "|")
            for row in rows[1:]:
                lines.append("| " + " | ".join(row) + " |")
            lines.append("")
            i += 1
        else:
            i += 1
    return "\n".join(lines) + "\n"


def render_txt(blocks: list[Block]) -> str:
    """Plain text: blocks separated by blank lines; bullets keep a '- ' prefix (the only
    structural cue a .txt file can carry); table rows are space-joined (no construct)."""
    lines = []
    for b in blocks:
        if b.kind in ("title", "narr", "raw"):
            lines.append(b.text)
            lines.append("")
        elif b.kind == "bullet":
            lines.append("- " + b.text)
            lines.append("")
        elif b.kind == "number":
            # keep a stable "1. " style prefix per item (reset handled loosely; plain text)
            lines.append("1. " + b.text)
            lines.append("")
        elif b.kind == "table":
            for row in b.extra:
                lines.append("  ".join(row))
            lines.append("")
    return "\n".join(lines) + "\n"


def render_docx(blocks: list[Block], path: Path, bullet_mode: str = "style") -> None:
    d = Document()
    for b in blocks:
        if b.kind == "title":
            d.add_heading(b.text, level=b.level)
        elif b.kind in ("narr", "raw"):
            d.add_paragraph(b.text)
        elif b.kind == "bullet":
            _add_docx_bullet(d, b.text, bullet_mode)
        elif b.kind == "number":
            d.add_paragraph(b.text, style="List Number")
        elif b.kind == "table":
            rows = b.extra
            t = d.add_table(rows=len(rows), cols=len(rows[0]))
            for r, row in enumerate(rows):
                for c, val in enumerate(row):
                    t.rows[r].cells[c].text = val
    d.save(str(path))


def _add_docx_bullet(d: Document, text: str, mode: str) -> None:
    if mode == "style":
        d.add_paragraph(text, style="List Bullet")
    elif mode == "dash":
        d.add_paragraph("- " + text)  # manual dash, Normal style
    elif mode == "numpr":
        p = d.add_paragraph(text)  # numbering props, no List* style
        pPr = p._p.get_or_add_pPr()
        numPr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl"); ilvl.set(qn("w:val"), "0"); numPr.append(ilvl)
        numId = OxmlElement("w:numId"); numId.set(qn("w:val"), "1"); numPr.append(numId)
        pPr.append(numPr)
    else:
        raise ValueError(mode)


# ---------------------------------------------------------------------------
# Logical documents
# ---------------------------------------------------------------------------
def doc_canonical() -> list[Block]:
    return [
        Block("d1-t1", "Title", "zztitle1", "Quarterly Operations Review zztitle1", "title", 1),
        Block("d1-n1", "NarrativeText", "zznarr1",
              "Our team zznarr1 evaluated fourteen scraping tools during the pilot and recorded every result carefully.", "narr"),
        Block("d1-n2", "NarrativeText", "zznarr2",
              "The engineers zznarr2 compared latency and memory across each candidate before writing the recommendation.", "narr"),
        Block("d1-t2", "Title", "zztitle2", "Key Findings zztitle2", "title", 2),
        Block("d1-b1", "ListItem", "zzitem1", "First zzitem1 the parser handled malformed markup without crashing", "bullet"),
        Block("d1-b2", "ListItem", "zzitem2", "Second zzitem2 it preserved the original list ordering throughout", "bullet"),
        Block("d1-b3", "ListItem", "zzitem3", "Third zzitem3 it exported clean readable plain text reliably", "bullet"),
        Block("d1-t3", "Title", "zztitle3", "Setup Steps zztitle3", "title", 2),
        Block("d1-o1", "ListItem", "zznum1", "Install zznum1 the project dependencies from the lockfile", "number"),
        Block("d1-o2", "ListItem", "zznum2", "Run zznum2 the extractor against the sample corpus", "number"),
        Block("d1-n3", "NarrativeText", "zznarr3",
              "In summary zznarr3 the tool performed consistently across all of the evaluated document formats.", "narr"),
    ]


def doc_table() -> list[Block]:
    return [
        Block("d2-t1", "Title", "zztbltitle", "Benchmark Results zztbltitle", "title", 1),
        Block("d2-n1", "NarrativeText", "zztblnarr",
              "The table zztblnarr below summarizes the measured throughput for every tool we tested.", "narr"),
        Block("d2-tab", "Table", "zztblcell",
              "table with sentinel cell zztblcell", "table",
              extra=[["Tool", "Throughput"], ["zztblcell alpha", "120"], ["beta engine", "95"]],
              note="txt has no table construct; expect NOT Table there"),
    ]


def doc_adversarial() -> list[Block]:
    return [
        Block("d3-over", "Title", "zzover12",
              "This zzover12 heading deliberately contains thirteen separate ordinary words exceeding the configured limit", "title", 1,
              note="13 words > title_max_word_length=12; structural h1/Heading should still be Title"),
        Block("d3-caps", "Title", "zzcaps",
              "SYSTEM ZZCAPS CONFIGURATION GUIDE", "title", 2,
              note="all-caps heading; cap ratio > 0.5"),
        Block("d3-verbless", "NarrativeText", "zzverbless",
              "A zzverbless comprehensive catalogue of vendor names, product categories, regional offices and contact records", "narr",
              note="no finite verb; is_possible_narrative_text should fail -> UncategorizedText in ALL carriers"),
        Block("d3-nonalpha", "UncategorizedText", "zznonalpha",
              "zznonalpha 12.5% :: 40/60 >> $$$ 3.14 ~~ 900", "raw",
              note="high non-alpha ratio > 0.5; UncategorizedText is the DEFENSIBLE output "
                   "(not scored as a NarrativeText miss) — kept as a consistency example"),
        Block("d3-realnarr", "NarrativeText", "zzrealnarr",
              "Finally zzrealnarr the reviewer summarized the outcome and approved the migration to the new pipeline.", "narr",
              note="control: a real verb-bearing sentence, expect NarrativeText everywhere"),
    ]


def doc_shortlist() -> list[Block]:
    """The #3280 shape: an intro PARAGRAPH immediately followed by a short-item list.
    The list is rendered md WITH a blank line before it (proper) AND WITHOUT (the lazy-
    continuation case), to isolate the Python-Markdown block rule. The block before the
    list is a paragraph on purpose — a heading always forces its own trailing blank, so
    only a paragraph-adjacent list exposes the collapse."""
    return [
        Block("d4-p", "NarrativeText", "zzsl_intro",
              "The following zzsl_intro items were recorded during the short test run", "narr"),
        Block("d4-b1", "ListItem", "zzsl1", "zzsl1 one", "bullet"),
        Block("d4-b2", "ListItem", "zzsl2", "zzsl2 two", "bullet"),
        Block("d4-b3", "ListItem", "zzsl3", "zzsl3 three", "bullet"),
        Block("d4-b4", "ListItem", "zzsl4", "zzsl4 four", "bullet"),
        Block("d4-b5", "ListItem", "zzsl5", "zzsl5 five", "bullet"),
    ]


def render_md_lazy_list(blocks: list[Block]) -> str:
    """Reproduce the #3280 lazy-continuation collapse: the first paragraph is placed on
    the line IMMEDIATELY above the bullet lines with NO blank separator, so Python-Markdown
    folds the whole thing into one <p>. (render_md always inserts a blank after each block,
    which markdown headings force anyway — hence this dedicated renderer.)"""
    lines = []
    intro = blocks[0]
    lines.append(intro.text)  # no trailing blank line -> list is a lazy continuation
    for b in blocks[1:]:
        lines.append("- " + b.text)
    return "\n".join(lines) + "\n"


def doc_large(reps: int = 60) -> list[Block]:
    """Resource-cost fixture: canonical blocks repeated to a sizeable document."""
    base = doc_canonical()
    blocks = []
    for r in range(reps):
        for b in base:
            # zero-pad + 'z' delimiter so sentinels are PREFIX-FREE (…z001 is not a
            # substring of …z010): avoids false SPLIT/MERGE matches in metrics.
            sent = f"{b.sentinel}z{r:03d}"
            txt = b.text.replace(b.sentinel, sent)
            extra = b.extra
            blocks.append(Block(f"{b.id}-r{r}", b.intent, sent, txt, b.kind, b.level, extra))
    return blocks


# ---------------------------------------------------------------------------
# Emit
# ---------------------------------------------------------------------------
def emit(name, blocks, formats=("html", "md", "txt", "docx"), md_blank=True, docx_bullet="style"):
    if "html" in formats:
        (FIX / f"{name}.html").write_text(render_html(blocks), encoding="utf-8")
    if "md" in formats:
        (FIX / f"{name}.md").write_text(render_md(blocks, blank_before_list=md_blank), encoding="utf-8")
    if "txt" in formats:
        (FIX / f"{name}.txt").write_text(render_txt(blocks), encoding="utf-8")
    if "docx" in formats:
        render_docx(blocks, FIX / f"{name}.docx", bullet_mode=docx_bullet)
    return {
        "name": name,
        "formats": list(formats),
        "blocks": [gt_of(b) for b in blocks],
    }


def main() -> int:
    docs = {}

    # D1 canonical — all four carriers, the fair baseline
    docs["d1_canonical"] = emit("d1_canonical", doc_canonical())

    # D2 table — html/md/docx have a table construct; txt does not
    docs["d2_table"] = emit("d2_table", doc_table())

    # D3 adversarial heuristic boundaries — all four carriers
    docs["d3_adversarial"] = emit("d3_adversarial", doc_adversarial())

    # D4 short list — md WITH blank line (proper) and WITHOUT (lazy-continuation collapse),
    # + html + txt. The blank variant is the fair reference; the lazy variant reproduces
    # the #3280 Python-Markdown fold.
    sl = doc_shortlist()
    docs["d4_shortlist_mdblank"] = emit("d4_shortlist_mdblank", sl, formats=("md", "html", "txt"), md_blank=True)
    (FIX / "d4_shortlist_mdlazy.md").write_text(render_md_lazy_list(sl), encoding="utf-8")
    docs["d4_shortlist_mdlazy"] = {"name": "d4_shortlist_mdlazy", "formats": ["md"],
                                   "blocks": [gt_of(b) for b in sl]}

    # D5 docx bullet authoring variants — docx only, 3 modes
    bl = doc_canonical()  # reuse; only the bullet blocks matter, but keep a full doc
    for mode in ("style", "dash", "numpr"):
        render_docx(bl, FIX / f"d5_docx_{mode}.docx", bullet_mode=mode)
        docs[f"d5_docx_{mode}"] = {"name": f"d5_docx_{mode}", "formats": ["docx"],
                                   "blocks": [gt_of(b) for b in bl],
                                   "docx_bullet_mode": mode}

    # D7 large — resource cost (html/md/txt/docx)
    docs["d7_large"] = emit("d7_large", doc_large(60))

    (FIX / "ground_truth.json").write_text(
        json.dumps(docs, indent=2, ensure_ascii=False) + "\n", encoding="utf-8"
    )
    n_docs = len(docs)
    n_blocks = sum(len(d["blocks"]) for d in docs.values())
    print(f"wrote {n_docs} fixture docs, {n_blocks} labeled block-instances -> {FIX}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
